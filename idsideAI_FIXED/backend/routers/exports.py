from fastapi import APIRouter, Query
from fastapi.responses import JSONResponse
from starlette.responses import FileResponse
from starlette.background import BackgroundTask
from starlette.staticfiles import StaticFiles
import os, io, json, time, zipfile
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from docx import Document

from backend.routers.graphs import _materialize_subgraph, _load_snapshots

router = APIRouter(prefix="/graphs", tags=["exports"])

EXPORT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)

def _load_snapshot_for(graph_id: str, sid: str):
    data = _load_snapshots(graph_id)
    for s in data.get("snapshots", []):
        if s["id"] == sid:
            return s
    return None

def _nmap(nodes):
    return {n["id"]: n for n in nodes}

def _emap(edges):
    return {e["id"]: e for e in edges}

def _derive_why(graph):
    nodes = graph["nodes"]; edges = graph["edges"]
    nm = _nmap(nodes)
    chosen_edge = next((e for e in edges if e.get("type","").upper()=="CHOSEN" or e.get("chosen")==True), None)
    decision = next((n for n in nodes if n.get("type")=="Decision"), None)
    chosen_option = nm.get(chosen_edge["target"]) if chosen_edge else None
    # evidence supporting chosen
    supports = [e for e in edges if e.get("type") in ("supported_by","SUPPORTED_BY") and e.get("target") == (chosen_option["id"] if chosen_option else None)]
    ev_nodes = [nm.get(e["source"]) for e in supports if nm.get(e["source"])]
    # naive rationale
    title = f"Why: {chosen_option['title'] if chosen_option else 'Decision'}"
    body = "We selected {} based on supporting evidence and risk/constraint considerations.".format(chosen_option['title'] if chosen_option else 'the option')
    citations = [{"id": n["id"], "title": n.get("title","Evidence"), "confidence": n.get("confidence"), "source": n.get("source_url"), "captured_at": n.get("captured_at")} for n in ev_nodes]
    return {"title": title, "body": body, "citations": citations, "decision": decision, "option": chosen_option}

@router.post("/export/audit")
def export_audit(graph_id: str, frm: str = Query(...), to: str = Query(...)):
    data = _load_snapshots(graph_id)
    a = next((s for s in data.get("snapshots", []) if s["id"] == frm), None)
    b = next((s for s in data.get("snapshots", []) if s["id"] == to), None)
    if not a or not b:
        return JSONResponse({"error":"Snapshots not found"}, status_code=404)
    manifest = {"graph_id": graph_id, "from": frm, "to": to, "created_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())}
    fname = f"audit_{graph_id}_{frm}_{to}_{int(time.time())}.zip"
    fpath = os.path.join(EXPORT_DIR, fname)
    with zipfile.ZipFile(fpath, "w") as z:
        z.writestr("manifest.json", json.dumps(manifest, indent=2))
        z.writestr(f"snapshot_{frm}.json", json.dumps({"nodes":a.get("nodes",[]), "edges":a.get("edges",[])}, indent=2))
        z.writestr(f"snapshot_{to}.json", json.dumps({"nodes":b.get("nodes",[]), "edges":b.get("edges",[])}, indent=2))
    return {"url": f"/exports/{fname}"}

@router.post("/export/why")
def export_why(graph_id: str, decision: str = Query(...), to: str = Query(None)):
    # Use snapshot 'to' if provided; else materialize live
    if to:
        s = _load_snapshot_for(graph_id, to)
        if not s:
            return JSONResponse({"error":"Snapshot not found"}, status_code=404)
        graph = {"nodes": s.get("nodes",[]), "edges": s.get("edges",[])}
    else:
        graph = _materialize_subgraph(root=decision, depth=3)
    why = _derive_why(graph)

    # Generate PDF
    pdf_name = f"why_{graph_id}_{int(time.time())}.pdf"
    pdf_path = os.path.join(EXPORT_DIR, pdf_name)
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    elements = [Paragraph("IDECIDE — Why Narrative", styles["Title"]), Spacer(1, 10)]
    elements.append(Paragraph(why["title"], styles["Heading2"]))
    elements.append(Paragraph(why["body"], styles["Normal"])); elements.append(Spacer(1,8))
    elements.append(Paragraph("Citations", styles["Heading3"]))
    rows = [["ID","Title","Confidence","Source","Captured"]]
    for c in why["citations"]:
        rows.append([c.get("id",""), c.get("title",""), str(c.get("confidence","")), c.get("source","") or "-", c.get("captured_at","") or "-"])
    if len(rows)>1:
        table = Table(rows, colWidths=[40,200,60,160,60])
        table.setStyle(TableStyle([('GRID',(0,0),(-1,-1),0.5,colors.HexColor("#E6EAF0"))]))
        elements.append(table)
    doc.build(elements)

    # Generate DOCX
    docx_name = f"why_{graph_id}_{int(time.time())}.docx"
    docx_path = os.path.join(EXPORT_DIR, docx_name)
    d = Document()
    d.add_heading("IDECIDE — Why Narrative", 0)
    d.add_heading(why["title"], level=1)
    d.add_paragraph(why["body"])
    d.add_heading("Citations", level=2)
    for c in why["citations"]:
        d.add_paragraph(f"{c.get('id')}: {c.get('title')} (conf {c.get('confidence')}) — {c.get('source') or '-'}")
    d.save(docx_path)

    return {"pdf": f"/exports/{pdf_name}", "docx": f"/exports/{docx_name}"}
